#!/usr/bin/python3.6
# -*-coding:utf-8-*-
import os
import json
import csv
import time
import re
from docx import Document
from electron_microscope import elec_lines
from docx.shared import Inches

res_dict = dict()


class FileHandler(object):

    def __init__(self, file_path):
        self.file_path = file_path
        # self.files_name_path = [os.path.join(self.folder_path, file_name)
        #                         for file_name in file_names
        #                         for _,_,file_names in os.walk(self.folder_path)]
        self.attr_lst = ['pathology_number',  #病理号
                         'name',  #姓名
                         'gender',  # 性别
                         'age',  # 年龄
                         'home',  # 籍贯
                         'clinical_diagnosis',  # 临床诊断
                         'hospital',  # 送检医院
                         'inpatient_area',  # 病区
                         'admission_number',  # 住院号
                         'report_date',
                         'light_microscopy',  # 光镜检查结果
                         'impression',  # 印象
                         'IgG',
                         'IgA',
                         'IgM',
                         'C3',
                         'C1q',
                         'Fg',
                         'verifier',  #检验者
                         'elec_info'
                         ]
        self.attr_dict = dict()
        self.attr_dict.setdefault('report_date', self.file_time(self.file_path))
        self.flsorescence = ['IgG', 'IgA', 'IgM', 'C3', 'C1q', 'Fg']  #荧光

        for attr in self.attr_lst:
            self.attr_dict.setdefault(attr, ' ')

    # def parsing_file(self):
    #     light_sign = 0
    #     f = Document(self.file_path)
    #     # content = [para.text for para in f.paragraphs]
    #
    #     for para in f.paragraphs:
    #         light_sign += 1
    #         line_content = para.text
    #
    #         print(line_content)
    #         content_list = re.split(r"[：]", line_content)
    #
    #         if r'姓名' in line_content and r'送检医院' in line_content:
    #             # print(re.sub(r'\s+', '  ', line_content))
    #             self.attr_dict['name'] = content_list[1][:8].strip().replace(" ", "")
    #             self.attr_dict['hospital'] = content_list[2].strip()
    #         elif r'病区' in line_content and r'性别' in line_content:
    #             # print(content_list)
    #             self.attr_dict['gender'] = content_list[1].strip().split()[0]
    #             self.attr_dict['inpatient_area'] = content_list[2].strip()
    #
    #         elif r'年龄' in line_content and r'住院号' in line_content:
    #             self.attr_dict['age'] = content_list[1].strip().split()[0]
    #             self.attr_dict['admission_number'] = content_list[2].strip()
    #
    #         elif r'籍贯' in line_content and '病理号' in line_content:
    #             self.attr_dict['home'] = content_list[1][:4]
    #             self.attr_dict['pathology_number'] = content_list[2].strip()
    #         elif r'临床诊断' in line_content:
    #             self.attr_dict['clinical_diagnosis'] = content_list[1].strip()
    #         elif content_list[0] in self.flsorescence:
    #             self.attr_dict[content_list[0]] = content_list[1].strip()
    #         elif r'光镜检查结果' in line_content:
    #             light_sign = -2
    #         elif light_sign == -1:
    #             self.attr_dict['light_microscopy'] = content_list[0]
    #         elif light_sign == 0:
    #             self.attr_dict['light_microscopy'] += content_list[0]
    #             self.attr_dict['light_microscopy'] = self.attr_dict['light_microscopy'].strip()
    #
    #         elif r'印象' in line_content:
    #             self.attr_dict['impression'] = content_list[1].strip()
    #         elif r'检验者' in line_content:
    #             self.attr_dict['verifier'] = content_list[1].strip()
    #         elif r'年' in line_content:
    #             self.attr_dict['report_date'] = content_list[0].strip()

    def parsing2_file(self):
        f = Document(self.file_path)
        # content = [para.text for para in f.paragraphs]
        info_lst = list()
        for para in f.paragraphs:
            line_content = para.text.strip()
            content_list = re.split(r"[：]", line_content)
            # content_list = re.split(r"[：]", line_content)
            if line_content and line_content != r"肾脏组织活检病理报告" \
                    and r"中山医科大学附属第一医院" not in line_content:

                if r'姓名' in line_content and r'送检医院' in line_content:
                    hospital_index = line_content.index(r'送检医院')
                    self.attr_dict['name'] = line_content[3:hospital_index].strip()
                    self.attr_dict['hospital'] = line_content[hospital_index+5:].strip()
                    continue
                elif r'性别' in line_content and r'病区' in line_content:
                    area_index = line_content.index(r'病区')
                    self.attr_dict['gender'] = line_content[3:area_index].strip()
                    self.attr_dict['inpatient_area'] = line_content[area_index+3:].strip()
                    continue
                elif r'年龄' in line_content and r'住院号' in line_content:
                    admission_number_index = line_content.index('住院号')
                    self.attr_dict['age'] = line_content[3:admission_number_index].strip()
                    self.attr_dict['admission_number'] = line_content[admission_number_index+4:].strip()
                    continue
                elif r'籍贯' in line_content and '病理号' in line_content:
                    patho_index = line_content.index(r'病理号：')
                    self.attr_dict['home'] = line_content[3:patho_index].strip()
                    self.attr_dict['pathology_number'] = line_content[patho_index+4:].strip()
                    continue
                elif r'临床诊断' in line_content:
                    self.attr_dict['clinical_diagnosis'] = content_list[1].strip()
                    continue
                elif content_list[0] in self.flsorescence:
                    self.attr_dict[content_list[0]] = content_list[1].strip()
                    continue
                # elif r'光镜检查结果' in line_content:
                #     light_sign = -2
                # elif light_sign == -1:
                #     self.attr_dict['light_microscopy'] = content_list[0]
                # elif light_sign == 0:
                #     self.attr_dict['light_microscopy'] += content_list[0]
                #     self.attr_dict['light_microscopy'] = self.attr_dict['light_microscopy'].strip()
                #
                # elif r'印象' in line_content:
                #     self.attr_dict['impression'] = content_list[1].strip()
                elif r'检验者' in line_content:
                    self.attr_dict['verifier'] = content_list[1].strip()
                    continue
                elif r'年' in line_content:
                    if r'日期' in line_content:
                        self.attr_dict['report_date'] = content_list[1].strip()
                    else:
                        self.attr_dict['report_date'] = content_list[0].strip()
                    continue
                info_lst.append(line_content.strip())
        self.special_handle(info_lst)

    def special_handle(self, info_lst):
        lightmicroscope_index = 0
        impression_index = 0
        # print(info_lst)
        for index_, info in enumerate(info_lst):
            if r'光镜检查结果' in info:
                lightmicroscope_index = index_
            elif r'印象：' in info:
                impression_index = index_
                break

        if not impression_index:
            impression_index = lightmicroscope_index + 2

        light_info_lst = info_lst[lightmicroscope_index+1: impression_index]
        light_info = '.'.join(light_info_lst)
        self.attr_dict['light_microscopy'] = light_info

        impression_info_lst = info_lst[impression_index:]
        impression_info = '.'.join(impression_info_lst)
        impression_info_index = impression_info.index("：")+1 if r'印象'in impression_info else 0
        self.attr_dict['impression'] = impression_info[impression_info_index:].strip()

        print(self.attr_dict)

    def line_creation(self):
        line = [self.attr_dict[at] for at in self.attr_lst]
        out = open('final2.csv', 'a', newline='')
        csv_write = csv.writer(out, dialect='excel')
        csv_write.writerow(line)
        out.close()
        pass

    def file_time(self, file_path):
        file_time_stamp = os.path.getctime(file_path)
        style_time = time.strftime("%Y年%m月%d日", time.localtime(file_time_stamp))
        return style_time

    def json_write(self):
        json_file = "2.json"
        with open(json_file, 'r', encoding='utf-8') as fw:
            load_dict = json.load(fw)
        load_dict["%s，%s" % (self.attr_dict['name'], self.attr_dict['admission_number'])] = \
            self.attr_dict['pathology_number']

        with open(json_file, 'w', encoding='utf-8') as dump_f:
            json.dump(load_dict, dump_f, ensure_ascii=False)

    def check_repe(self, elec_):
        light_name_number = "%s,%s" % (self.attr_dict['name'], self.attr_dict['admission_number'])
        if light_name_number in elec_:
            self.attr_dict['elec_info'] = elec_[light_name_number][-1]
        # TODO: also need to extract unused data to another file

        file = open("repeat_record", encoding="utf-8", mode="a")
        file.write(light_name_number)
        file.close()


        pass


def traverse_all_files(folder_path):
    electron_info = elec_lines()
    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            file_path = os.path.join(folder_path, filename)
            print(file_path)
            fh = FileHandler(file_path)
            fh.parsing2_file()
            fh.check_repe(electron_info)
            fh.line_creation()
            fh.json_write()
    # for dirpath, dirnames, filenames in os.walk(self.folder_path):
    #     for file_name in filenames:
    #         file_path = os.path.join(dirpath,file_name)


if __name__ == '__main__':
    folder = '/'
    traverse_all_files(folder)


